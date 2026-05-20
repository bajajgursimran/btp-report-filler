import io
import json
import os
import re
import zipfile
import html as html_lib
from typing import List
import xml.etree.ElementTree as ET

import anthropic
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from docx import Document

app = FastAPI(title="SAP ICC Certificate Report Generation Tool")

# ── static files ──────────────────────────────────────────────────────────────
@app.get("/")
async def serve_index():
    return FileResponse("static/index.html")

app.mount("/static", StaticFiles(directory="static"), name="static")

# ── constants ─────────────────────────────────────────────────────────────────
FIELD_PAT = re.compile(
    r'(<w:fldChar w:fldCharType="begin"><w:ffData>.*?</w:ffData></w:fldChar>'
    r'.*?<w:fldChar w:fldCharType="separate"/>)'
    r'(.*?)'
    r'(<w:fldChar w:fldCharType="end"/>)',
    re.DOTALL,
)

RPR_DEFAULT = (
    '<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
    '<w:sz w:val="20"/></w:rPr>'
)

def _esc(t: str) -> str:
    return html_lib.escape(str(t), quote=False)

# ── source extraction ─────────────────────────────────────────────────────────
SOURCE_FIELD_PAT = re.compile(
    r'(<w:fldChar w:fldCharType="begin"><w:ffData>.*?</w:ffData></w:fldChar>'
    r'.*?<w:fldChar w:fldCharType="separate"/>)'
    r'(.*?)'
    r'(<w:fldChar w:fldCharType="end"/>)',
    re.DOTALL,
)

def _is_checked_in_source(ffdata_block: str) -> bool:
    """Return True if a source checkbox is checked — handles both <w:checked/> and <w:default w:val="1"/>."""
    if '<w:checked/>' in ffdata_block:
        return True
    m = re.search(r'<w:default w:val="(\d+)"', ffdata_block)
    if m and m.group(1) == "1":
        return True
    return False


def extract_source_text(source_bytes_list: list[bytes]) -> str:
    """
    Extract all readable text from source .docx/.docm files into a single plain-text
    string that Claude can read.
    - Checkboxes: emits "[CHECKED] label" or "[UNCHECKED] label" lines
    - Tables: label: value pairs
    - Paragraphs: plain text
    """
    sections: list[str] = []

    for raw in source_bytes_list:
        buf = io.BytesIO(raw)

        # get raw XML for checkbox extraction
        try:
            with zipfile.ZipFile(buf, 'r') as z:
                doc_xml = z.read('word/document.xml').decode('utf-8')
        except Exception:
            continue

        # build checkbox lines from raw XML
        checkbox_lines: list[str] = []
        for m in SOURCE_FIELD_PAT.finditer(doc_xml):
            g1 = m.group(1)
            if '<w:checkBox>' not in g1:
                continue
            checked = _is_checked_in_source(g1)
            after = doc_xml[m.end(): m.end() + 600]
            stop = re.search(r'<w:fldChar|</w:tc>|<w:tc[ >]', after)
            if stop:
                after = after[:stop.start()]
            texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', after, re.DOTALL)
            label = ' '.join(t.strip() for t in texts if t.strip())
            if label:
                checkbox_lines.append(f"[{'CHECKED' if checked else 'UNCHECKED'}] {label}")

        if checkbox_lines:
            sections.append("Checkbox states:\n" + "\n".join(checkbox_lines))

        # also extract table text via python-docx for non-checkbox fields
        buf.seek(0)
        try:
            doc = Document(buf)
        except Exception:
            buf.seek(0)
            plain_buf = io.BytesIO()
            with zipfile.ZipFile(buf, 'r') as zin, \
                 zipfile.ZipFile(plain_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == '[Content_Types].xml':
                        data = data.replace(
                            b'application/vnd.ms-word.document.macroEnabled.main+xml',
                            b'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
                        )
                    zout.writestr(item, data)
            plain_buf.seek(0)
            doc = Document(plain_buf)

        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paras:
            sections.append("\n".join(paras))

        for table in doc.tables:
            rows_text: list[str] = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0] == cells[1]:
                    rows_text.append(f"\n[Section: {cells[0]}]")
                    continue
                if len(cells) >= 2 and cells[0] and cells[1]:
                    value = cells[1].replace("\n", "; ")
                    rows_text.append(f"{cells[0]}: {value}")
                elif len(cells) >= 1 and cells[0]:
                    rows_text.append(cells[0])
            if rows_text:
                sections.append("\n".join(rows_text))

    return "\n\n---\n\n".join(sections)


# ── template field catalogue ──────────────────────────────────────────────────
def extract_field_catalogue(doc_xml: str) -> list[dict]:
    """
    Walk all form fields and numbered test-case rows and return a list of
    field descriptors that Claude will fill in.
    """
    fields: list[dict] = []
    seen: dict[str, int] = {}

    for m in FIELD_PAT.finditer(doc_xml):
        g1 = m.group(1)
        is_checkbox = '<w:checkBox>' in g1
        is_dropdown = '<w:ddList>'   in g1

        nm    = re.search(r'<w:name w:val="([^"]+)"', g1)
        fname = nm.group(1) if nm else "unknown"
        idx   = seen.get(fname, 0)
        seen[fname] = idx + 1
        key   = f"{fname}[{idx}]"

        # context: visible text just before this field (section heading / row label)
        snippet      = doc_xml[max(0, m.start() - 1500): m.start()]
        ctx_texts    = re.findall(r'<w:t[^>]*>(.*?)</w:t>', snippet, re.DOTALL)
        ctx_words    = [t.strip() for t in ctx_texts if t.strip() and '<w' not in t]
        context      = " ".join(ctx_words[-15:])

        # for text fields: also use the ffData default value as context hint
        default_hint = ""
        dfl = re.search(r'<w:default w:val="([^"]+)"', g1)
        if dfl:
            default_hint = html_lib.unescape(dfl.group(1))

        # label: text immediately AFTER this field up to the next field/cell boundary
        after_snippet = doc_xml[m.end(): m.end() + 600]
        stop = re.search(r'<w:fldChar|</w:tc>|<w:tc[ >]', after_snippet)
        if stop:
            after_snippet = after_snippet[:stop.start()]
        after_texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', after_snippet, re.DOTALL)
        after_words = [t.strip() for t in after_texts if t.strip()]
        label       = " ".join(after_words) if after_words else ""

        if is_checkbox:
            already = '<w:checked/>' in g1
            fields.append({"key": key, "type": "checkbox",
                           "context": context, "label": label, "current": already})
        elif is_dropdown:
            entries = re.findall(r'<w:listEntry w:val="([^"]+)"', g1)
            fields.append({"key": key, "type": "dropdown",
                           "context": context, "options": entries})
        else:
            cur_text = " ".join(re.findall(r'<w:t[^>]*>(.*?)</w:t>',
                                           m.group(2), re.DOTALL)).strip()
            # build rich context: surrounding text + default hint
            rich_context = context
            if default_hint and default_hint not in context:
                rich_context = f"{context} [field hint: {default_hint}]".strip()
            fields.append({"key": key, "type": "text",
                           "context": rich_context, "current": cur_text})

    # numbered test-case rows (plain table cells, not form fields)
    # Build a position->section map by finding section headings in document order
    section_markers: list[tuple[int, str]] = []
    for sm in re.finditer(r'<w:t[^>]*>(.*?)</w:t>', doc_xml, re.DOTALL):
        txt = sm.group(1).strip().lower()
        if "error handling" in txt or "error test" in txt:
            section_markers.append((sm.start(), "error_handling"))
        elif "functional test" in txt or "functional" in txt:
            section_markers.append((sm.start(), "functional"))

    def _section_for_pos(pos: int) -> str:
        """Return the most recent section heading before pos."""
        result = "functional"
        for marker_pos, marker_section in section_markers:
            if marker_pos < pos:
                result = marker_section
            else:
                break
        return result

    test_row_seen: dict[str, int] = {}
    for m in re.finditer(r'<w:t>(\d+\.)</w:t></w:r>', doc_xml):
        snippet   = doc_xml[max(0, m.start() - 400): m.start()]
        ctx_texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', snippet, re.DOTALL)
        ctx_words = [t.strip() for t in ctx_texts if t.strip()]
        context   = " ".join(ctx_words[-6:])
        section   = _section_for_pos(m.start())
        base_key  = f"test_row_{section}_{m.group(1)}"
        idx = test_row_seen.get(base_key, 0)
        test_row_seen[base_key] = idx + 1
        key = base_key if idx == 0 else f"{base_key}_{idx}"
        fields.append({"key": key,
                       "type": "test_row",
                       "label": m.group(1),
                       "section": section,
                       "context": context})

    # plain empty table cells (label | empty value, no form field)
    rows = re.findall(r'<w:tr[ >].*?</w:tr>', doc_xml, re.DOTALL)
    empty_cell_seen: dict[str, int] = {}
    for row in rows:
        cells = re.findall(r'<w:tc>.*?</w:tc>', row, re.DOTALL)
        if len(cells) < 2:
            continue
        left_texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', cells[0])
        right_texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', cells[1])
        right_has_field = '<w:fldChar' in cells[1]
        label = ' '.join(t.strip() for t in left_texts if t.strip() and '<w' not in t)
        right_content = ' '.join(t.strip() for t in right_texts if t.strip())
        if label and not right_content and not right_has_field:
            idx = empty_cell_seen.get(label, 0)
            empty_cell_seen[label] = idx + 1
            key = f"plain_cell_{re.sub(r'[^a-zA-Z0-9]', '_', label[:30])}[{idx}]"
            fields.append({"key": key, "type": "plain_cell", "context": label})

    return fields


# ── Claude AI fill ────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are an expert SAP BTP certification consultant filling in a
BTP Certificate Test Report (Word .docm template).

You will receive:
1. SOURCE DATA — all text extracted from the partner's source documents (e.g. Technical Product Profile)
2. FIELD CATALOGUE — a JSON list of every field in the template, with its type and surrounding context

Your job: return a JSON object mapping each field "key" to the value it should receive.

Rules:
- "text" and "plain_cell" fields → return a string value (or null to leave blank).
  Use the "context" and "field hint" to understand what information belongs in the field.
  Match semantically — the field label in the template may be worded differently from the source data label, but mean the same thing.
  For example: "Overview of the purpose and functionality" should be filled with the FULL multi-sentence description of what the application does and its purpose — never just the product name or a single word. Look for source fields like "Please give a broad overview on the functionality and the purpose", "Description", "Summary", or any paragraph that explains what the product does.
  "Technical approach" should be filled from architecture/technology description sections in the source.
  Always prefer the full descriptive answer from the source over a short name or version number.
  If a field is about "overview", "purpose", "functionality", "description", or "summary" — always return at least 2-3 sentences explaining WHAT the product does and WHY, not just its name.
- "checkbox" fields → each checkbox has a "label" (the option it represents) and a "context" (the question/section it belongs to).
  Determine whether to check it using ALL available evidence in priority order:
  1. If the source data has a "Checkbox states" section with "[CHECKED] label" or "[UNCHECKED] label" — use that as ground truth.
  2. If no checkbox states section exists, look for the label text appearing as a selected/confirmed value in table cells or paragraph text.
  3. If the source data lists options as plain text in a cell (e.g. "Destination service; Cloud Connector"), treat ALL listed items as selected.
  Return true to check, false to uncheck, null if genuinely unclear.
  IMPORTANT: Be conservative — only check boxes clearly confirmed by source data.
  For mutually exclusive groups, check ONLY the one(s) explicitly stated in the source.
  Some labels contain placeholders like "<mention the version>" — match on the base product name only.
- "dropdown" fields → return the string exactly matching one of the listed options, or null
- "test_row" fields → return a full test case string (multi-line is fine, use \\n)
  - Fields with "section": "functional" → derive from the partner's main functional capabilities and business processes
  - Fields with "section": "error_handling" → derive from error handling, exception flows, and failure scenarios described in the source
  - These two sections MUST be different — do not reuse functional test cases for error handling rows
  - Include: test case name, description, steps, expected result
- Only use information present in the source data — do not invent partner details
- For fields where source data has no relevant information, return null
- Keep values concise and professional
- For certification date / expiry / cert ID fields: return null (ICC fills these)
- For ICC consultant name: return null (ICC fills this)

Return ONLY valid JSON. No explanation, no markdown fences."""

def ai_fill(source_text: str, field_catalogue: list[dict],
            api_key: str) -> dict:
    """
    Ask Claude to map every template field to a value from the source data.
    Returns {field_key: value_or_null}.
    Uses prompt caching on the large source_text block.
    """
    base_url = os.environ.get("ANTHROPIC_BASE_URL")
    auth_token = os.environ.get("ANTHROPIC_AUTH_TOKEN")

    if base_url and auth_token:
        client = anthropic.Anthropic(auth_token=auth_token, base_url=base_url)
    else:
        client = anthropic.Anthropic(api_key=api_key)

    user_content = (
        "SOURCE DATA:\n"
        + source_text
        + "\n\n---\n\nFIELD CATALOGUE:\n"
        + json.dumps(field_catalogue, indent=2)
        + "\n\nReturn the JSON mapping now."
    )

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        system=[
            {
                "type": "text",
                "text": SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},   # cache system prompt
            }
        ],
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": user_content,
                        "cache_control": {"type": "ephemeral"},  # cache source data
                    }
                ],
            }
        ],
    )

    raw = response.content[0].text.strip()
    # strip accidental markdown fences
    raw = re.sub(r'^```[a-z]*\n?', '', raw)
    raw = re.sub(r'\n?```$', '', raw)
    return json.loads(raw)


# ── XML injection helpers ─────────────────────────────────────────────────────
def _make_runs(text: str, rpr: str) -> str:
    parts = []
    for i, line in enumerate(text.split("\n")):
        if i:
            parts.append(f'<w:r>{rpr}<w:br/></w:r>')
        parts.append(
            f'<w:r>{rpr}<w:t xml:space="preserve">{_esc(line)}</w:t></w:r>'
        )
    return "".join(parts)


def apply_ai_values(doc_xml: str, ai_values: dict) -> tuple[str, dict]:
    """
    Inject AI-resolved values into the template XML.
    Returns (filled_xml, summary).
    """
    summary: dict[str, str] = {}
    seen: dict[str, int] = {}

    def _replace(m: re.Match) -> str:
        g1, g2, g3 = m.group(1), m.group(2), m.group(3)

        is_checkbox = '<w:checkBox>' in g1
        is_dropdown = '<w:ddList>'   in g1

        nm    = re.search(r'<w:name w:val="([^"]+)"', g1)
        fname = nm.group(1) if nm else "unknown"
        idx   = seen.get(fname, 0)
        seen[fname] = idx + 1
        key   = f"{fname}[{idx}]"

        value = ai_values.get(key)

        # ── FORMCHECKBOX ──────────────────────────────────────────────────────
        if is_checkbox:
            if value is None:
                summary[key] = "unchanged"
                return m.group(0)
            already = '<w:checked/>' in g1
            if value is True and not already:
                new_g1 = g1.replace('<w:checkBox>', '<w:checkBox><w:checked/>', 1)
                summary[key] = "checked"
                return new_g1 + g2 + g3
            elif value is False and already:
                new_g1 = g1.replace('<w:checked/>', '')
                summary[key] = "unchecked"
                return new_g1 + g2 + g3
            summary[key] = "unchanged"
            return m.group(0)

        # ── FORMDROPDOWN ─────────────────────────────────────────────────────
        if is_dropdown:
            entries = re.findall(r'<w:listEntry w:val="([^"]+)"', g1)
            if value and value in entries:
                result_idx = entries.index(value)
            else:
                result_idx = 0   # default to first entry (OK)
            if '<w:result' not in g1:
                new_g1 = g1.replace(
                    '</w:ddList>',
                    f'</w:ddList><w:result w:val="{result_idx}"/>', 1
                )
                summary[key] = f"set:{value or entries[0]}"
                return new_g1 + g2 + g3
            summary[key] = "unchanged"
            return m.group(0)

        # ── FORMTEXT ─────────────────────────────────────────────────────────
        if not value:
            summary[key] = "no_match"
            return m.group(0)

        rpr_m = re.search(r'(<w:rPr>.*?</w:rPr>)', g2, re.DOTALL)
        rpr   = rpr_m.group(1) if rpr_m else RPR_DEFAULT

        new_g2 = (
            f'</w:r>'
            f'<w:r>{rpr}<w:t xml:space="preserve">{_esc(str(value))}</w:t></w:r>'
            f'<w:r>{rpr}'
        )
        summary[key] = f"filled:{str(value)[:40]}"
        return g1 + new_g2 + g3

    filled = FIELD_PAT.sub(_replace, doc_xml)

    # ── numbered test-case rows ───────────────────────────────────────────────
    # Build position->section map same as in extract_field_catalogue
    section_markers2: list[tuple[int, str]] = []
    for sm in re.finditer(r'<w:t[^>]*>(.*?)</w:t>', filled, re.DOTALL):
        txt = sm.group(1).strip().lower()
        if "error handling" in txt or "error test" in txt:
            section_markers2.append((sm.start(), "error_handling"))
        elif "functional test" in txt or "functional" in txt:
            section_markers2.append((sm.start(), "functional"))

    def _section_for_pos2(pos: int) -> str:
        result = "functional"
        for marker_pos, marker_section in section_markers2:
            if marker_pos < pos:
                result = marker_section
            else:
                break
        return result

    label_hits = list(re.finditer(r'<w:t>(\d+\.)</w:t></w:r>', filled))
    assignments: list[tuple[re.Match, str]] = []
    test_row_seen2: dict[str, int] = {}
    for hit in label_hits:
        section  = _section_for_pos2(hit.start())
        base_key = f"test_row_{section}_{hit.group(1)}"
        idx = test_row_seen2.get(base_key, 0)
        test_row_seen2[base_key] = idx + 1
        row_key = base_key if idx == 0 else f"{base_key}_{idx}"
        val = ai_values.get(row_key)
        if val:
            assignments.append((hit, str(val)))

    for hit, case_text in reversed(assignments):
        search_str = f'<w:t>{hit.group(1)}</w:t></w:r>'
        rpr_m = re.search(
            r'(<w:rPr>.*?</w:rPr>)<w:t>' + re.escape(hit.group(1)) + r'</w:t>',
            filled[max(0, hit.start() - 300): hit.end()], re.DOTALL,
        )
        rpr = rpr_m.group(1) if rpr_m else RPR_DEFAULT
        replacement = search_str + _make_runs(case_text, rpr)
        idx = filled.find(search_str, hit.start() - 10)
        if idx != -1:
            filled = filled[:idx] + replacement + filled[idx + len(search_str):]
            summary[f"test_row_{hit.group(1)}"] = f"filled:{case_text[:40]}"

    # ── plain empty table cells ───────────────────────────────────────────────
    plain_cell_seen: dict[str, int] = {}
    rows = list(re.finditer(r'<w:tr[ >].*?</w:tr>', filled, re.DOTALL))
    offset = 0
    for row_m in rows:
        row_xml = filled[row_m.start() + offset: row_m.end() + offset]
        cells = list(re.finditer(r'<w:tc>.*?</w:tc>', row_xml, re.DOTALL))
        if len(cells) < 2:
            continue
        left_texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', cells[0].group(0))
        right_texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', cells[1].group(0))
        right_has_field = '<w:fldChar' in cells[1].group(0)
        label = ' '.join(t.strip() for t in left_texts if t.strip() and '<w' not in t)
        right_content = ' '.join(t.strip() for t in right_texts if t.strip())
        if not label or right_content or right_has_field:
            continue
        cell_idx = plain_cell_seen.get(label, 0)
        plain_cell_seen[label] = cell_idx + 1
        key = f"plain_cell_{re.sub(r'[^a-zA-Z0-9]', '_', label[:30])}[{cell_idx}]"
        val = ai_values.get(key)
        if not val:
            continue
        # inject text into the empty right cell's first paragraph
        right_cell = cells[1].group(0)
        # find the empty paragraph and insert a run with the value
        rpr_m = re.search(r'(<w:rPr>.*?</w:rPr>)', row_xml, re.DOTALL)
        rpr = rpr_m.group(1) if rpr_m else RPR_DEFAULT
        new_run = _make_runs(str(val), rpr)
        new_right_cell = re.sub(
            r'(<w:p[^>]*>)(<w:pPr>.*?</w:pPr>)(</w:p>)',
            lambda pm: pm.group(1) + pm.group(2) + new_run + pm.group(3),
            right_cell, count=1, flags=re.DOTALL
        )
        if new_right_cell == right_cell:
            # no pPr — simpler paragraph
            new_right_cell = re.sub(
                r'(<w:p[^>]*>)(</w:p>)',
                lambda pm: pm.group(1) + new_run + pm.group(2),
                right_cell, count=1, flags=re.DOTALL
            )
        abs_start = row_m.start() + offset + cells[1].start()
        abs_end   = row_m.start() + offset + cells[1].end()
        filled = filled[:abs_start] + new_right_cell + filled[abs_end:]
        offset += len(new_right_cell) - len(right_cell)
        summary[key] = f"filled:{str(val)[:40]}"

    return filled, summary


# ── docm rebuilder (VBA-safe) ─────────────────────────────────────────────────
def rebuild_docm(template_bytes: bytes, new_xml: str) -> bytes:
    out_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(template_bytes), 'r') as zin, \
         zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == 'word/document.xml':
                zout.writestr(item, new_xml.encode('utf-8'))
            else:
                zout.writestr(item, zin.read(item.filename))
    return out_buf.getvalue()


# ── API ───────────────────────────────────────────────────────────────────────
@app.post("/api/fill")
async def fill_report(
    template: UploadFile = File(...),
    sources: List[UploadFile] = File(...),
):
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    has_proxy = os.environ.get("ANTHROPIC_AUTH_TOKEN") and os.environ.get("ANTHROPIC_BASE_URL")
    if not api_key and not has_proxy:
        raise HTTPException(500, "ANTHROPIC_API_KEY environment variable not set")

    if not template.filename.lower().endswith('.docm'):
        raise HTTPException(400, "Template must be a .docm file")
    for s in sources:
        if not s.filename.lower().endswith(('.docx', '.docm', '.doc')):
            raise HTTPException(400, f"{s.filename} must be a .docx or .docm file")

    template_bytes = await template.read()
    source_bytes   = [await s.read() for s in sources]

    # read template XML
    try:
        with zipfile.ZipFile(io.BytesIO(template_bytes), 'r') as z:
            doc_xml = z.read('word/document.xml').decode('utf-8')
    except Exception as exc:
        raise HTTPException(422, f"Could not read template: {exc}")

    # extract source text
    try:
        source_text = extract_source_text(source_bytes)
    except Exception as exc:
        raise HTTPException(422, f"Could not read source files: {exc}")

    # build field catalogue
    field_catalogue = extract_field_catalogue(doc_xml)

    # ask Claude to fill
    try:
        ai_values = ai_fill(source_text, field_catalogue, api_key)
    except json.JSONDecodeError as exc:
        raise HTTPException(500, f"Claude returned invalid JSON: {exc}")
    except anthropic.APIError as exc:
        raise HTTPException(502, f"Claude API error: {exc}")
    except Exception as exc:
        raise HTTPException(500, f"AI fill error: {exc}")

    # inject into XML
    try:
        filled_xml, summary = apply_ai_values(doc_xml, ai_values)
    except Exception as exc:
        raise HTTPException(500, f"XML injection error: {exc}")

    # validate XML
    try:
        ET.fromstring(filled_xml)
    except ET.ParseError as exc:
        raise HTTPException(500, f"Filled document has invalid XML: {exc}")

    # rebuild docm
    try:
        filled_bytes = rebuild_docm(template_bytes, filled_xml)
    except Exception as exc:
        raise HTTPException(500, f"Could not rebuild docm: {exc}")

    filled_count  = sum(1 for v in summary.values()
                        if v.startswith("filled") or v in ("checked", "unchecked")
                        or v.startswith("set:"))
    skipped_count = sum(1 for v in summary.values()
                        if v in ("unchanged", "no_match"))

    output_name = template.filename.replace('.docm', '_filled.docm')

    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.ms-word.document.macroEnabled.12",
        headers={
            "Content-Disposition": f'attachment; filename="{output_name}"',
            "Access-Control-Expose-Headers": "X-Fill-Summary, Content-Disposition",
            "X-Fill-Summary": json.dumps({
                "filled":  filled_count,
                "skipped": skipped_count,
                "total":   len(summary),
                "details": summary,
            }),
        },
    )
